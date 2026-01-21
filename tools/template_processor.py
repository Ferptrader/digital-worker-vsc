"""
Template Processor - Processa templates Word/Markdown com marcadores [[...]]
Preenche marcadores com dados gerados pelos agentes e gera documentos finais
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import re
from typing import Dict, Any, List, Optional


class TemplateProcessor:
    """
    Processa templates com marcadores [[NOME_MARCADOR]] e gera documentos finais
    """
    
    def __init__(self, templates_path="templates"):
        self.templates_path = Path(templates_path)
    
    def load_template(self, template_name: str) -> str:
        """
        Carrega template Markdown
        
        Args:
            template_name: Nome do template (ex: 'TEMPLATE_QI_COM_MARCADORES.md')
        
        Returns:
            Conteúdo do template como string
        """
        template_file = self.templates_path / template_name
        
        if not template_file.exists():
            raise FileNotFoundError(f"Template não encontrado: {template_file}")
        
        with open(template_file, 'r', encoding='utf-8') as f:
            return f.read()
    
    def replace_placeholders(self, template_content: str, context: Dict[str, Any]) -> str:
        """
        Substitui todos os [[MARCADORES]] pelos valores do contexto
        
        Args:
            template_content: Conteúdo do template com marcadores
            context: Dicionário com valores {NOME_MARCADOR: valor}
        
        Returns:
            Template preenchido
        """
        filled_content = template_content
        
        # Substituir marcadores
        for key, value in context.items():
            placeholder = f"[[{key}]]"
            filled_content = filled_content.replace(placeholder, str(value))
        
        # Verificar marcadores não preenchidos
        remaining = re.findall(r'\[\[[A-Za-z0-9_]+\]\]', filled_content)
        if remaining:
            unique_remaining = set(remaining)
            print(f"⚠️  AVISO: {len(unique_remaining)} marcadores não preenchidos:")
            for marker in sorted(unique_remaining):
                print(f"   - {marker}")
        
        return filled_content
    
    def markdown_to_docx(self, markdown_content: str, output_path: str) -> None:
        """
        Converte Markdown para DOCX com formatação básica
        
        Args:
            markdown_content: Conteúdo markdown preenchido
            output_path: Caminho para salvar o .docx
        """
        doc = Document()
        
        # Configurar estilos padrão
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        lines = markdown_content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i]
            
            # Linha vazia
            if not line.strip():
                i += 1
                continue
            
            # Títulos
            if line.startswith('# '):
                heading = doc.add_heading(line[2:].strip(), level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif line.startswith('## '):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:].strip(), level=3)
            
            # Negrito forte
            elif line.startswith('**') and line.endswith('**'):
                para = doc.add_paragraph()
                run = para.add_run(line[2:-2])
                run.bold = True
                run.font.size = Pt(12)
            
            # Lista com marcadores
            elif line.startswith('- '):
                doc.add_paragraph(line[2:].strip(), style='List Bullet')
            
            # Lista numerada
            elif re.match(r'^\d+\.\s', line):
                content = re.sub(r'^\d+\.\s', '', line)
                doc.add_paragraph(content, style='List Number')
            
            # Tabela Markdown
            elif line.startswith('|'):
                table_lines = []
                while i < len(lines) and lines[i].startswith('|'):
                    table_lines.append(lines[i])
                    i += 1
                i -= 1  # Voltar 1 pois o loop principal vai incrementar
                
                if len(table_lines) > 2:  # Header + separator + data
                    self._add_table_to_doc(doc, table_lines)
            
            # Parágrafo normal
            else:
                doc.add_paragraph(line)
            
            i += 1
        
        doc.save(output_path)
        print(f"✅ Documento Word salvo: {output_path}")
    
    def _add_table_to_doc(self, doc: Document, table_lines: List[str]) -> None:
        """
        Adiciona tabela Markdown ao documento Word
        
        Args:
            doc: Documento Word
            table_lines: Linhas da tabela em Markdown
        """
        # Parsear células
        rows_data = []
        for line in table_lines:
            cells = [cell.strip() for cell in line.split('|')[1:-1]]  # Remove | inicial e final
            # Pular linha separadora (contém apenas -, :, |)
            if all(c in '-:| ' for c in line):
                continue
            rows_data.append(cells)
        
        if not rows_data:
            return
        
        # Criar tabela
        num_cols = len(rows_data[0])
        table = doc.add_table(rows=len(rows_data), cols=num_cols)
        table.style = 'Light Grid Accent 1'
        
        # Preencher células
        for row_idx, row_data in enumerate(rows_data):
            for col_idx, cell_value in enumerate(row_data):
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = cell_value
                
                # Header em negrito
                if row_idx == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True


def generate_document(tipo_documento: str, context: Dict[str, Any], output_path: str) -> str:
    """
    Função principal para gerar documentos a partir de templates
    
    Args:
        tipo_documento: Tipo do documento ('QI_ANEXOS', 'OQ_ANEXOS', 'PV_VSC', etc)
        context: Dicionário com todos os valores para preencher marcadores
        output_path: Caminho completo para salvar o documento final (.docx)
    
    Returns:
        Caminho do arquivo gerado
    
    Raises:
        ValueError: Se tipo de documento não for suportado
        FileNotFoundError: Se template não existir
    """
    processor = TemplateProcessor()
    
    # Mapear tipo de documento para template
    template_map = {
        'QI_ANEXOS': 'TEMPLATE_QI_COM_MARCADORES.md',
        'OQ_ANEXOS': 'TEMPLATE_OQ_COM_MARCADORES.md',
        'PQ_ANEXOS': 'TEMPLATE_PQ_COM_MARCADORES.md',
        'PV_VSC': 'TEMPLATE_PV_VSC_COM_MARCADORES.md',
        'ARSC': 'TEMPLATE_ARSC_COM_MARCADORES.md',
        # Adicionar outros templates conforme necessário
    }
    
    if tipo_documento not in template_map:
        raise ValueError(
            f"Tipo de documento '{tipo_documento}' não suportado. "
            f"Tipos disponíveis: {', '.join(template_map.keys())}"
        )
    
    template_name = template_map[tipo_documento]
    
    print(f"📄 Gerando documento: {tipo_documento}")
    print(f"📋 Template: {template_name}")
    print(f"💾 Saída: {output_path}")
    
    # 1. Carregar template
    template_content = processor.load_template(template_name)
    
    # 2. Substituir marcadores
    filled_content = processor.replace_placeholders(template_content, context)
    
    # 3. Gerar DOCX
    processor.markdown_to_docx(filled_content, output_path)
    
    return output_path


# Exemplo de uso (para teste)
if __name__ == "__main__":
    # Dados de exemplo para QI do SensorWeb
    exemplo_context = {
        'NOME_SISTEMA': 'SENSORWEB',
        'CODIGO_DOCUMENTO': 'PT-QISC-SENS-001',
        'VERSAO_DOCUMENTO': '01',
        'DATA_DOCUMENTO': '21/01/2026',
        'NOME_EMPRESA': 'KIVALITA',
        'NOME_PROTOCOLO': 'PT-QISC-SENS-001 - Qualificação de Instalação',
        
        # Teste TI-01
        'CORRIDA_TESTE_TI_01': '001',
        'RESULTADO_TESTE_TI_01_1_1': 'Aprovado',
        'RESULTADO_TESTE_TI_01_1_2': 'Aprovado',
        'EVIDENCIA_TESTE_TI_01_1_1': 'Anexo A-1',
        'OBSERVACOES_TESTE_TI_01': 'Todos os critérios atendidos',
        'RESULTADO_CONSOLIDADO_TI_01': 'Aprovado',
        'APROVADOR_TI_01': 'João Silva',
        'DATA_APROVACAO_TI_01': '21/01/2026',
        
        # Assinatura
        'FUNCAO_RESPONSAVEL_QI': 'Gerente de Validação',
        'NOME_RESPONSAVEL_QI': 'Maria Santos',
        'DATA_ASSINATURA_CONCLUSAO_QI': '21/01/2026',
    }
    
    # Gerar documento
    generate_document(
        tipo_documento='QI_ANEXOS',
        context=exemplo_context,
        output_path='outputs/PT-QISC-SENS-001_QI_ANEXOS_v01.docx'
    )
